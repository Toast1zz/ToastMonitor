from docx import Document
from docx.oxml.ns import qn

src = "/Users/toast1/Downloads/Corrected_Kilishop_Production_Proposal.docx"
out = "/Users/toast1/Projects/ToastMonitor/Corrected_Kilishop_Production_Proposal_中文翻译.docx"

paragraphs = {
0: "KILISHOP 内容制作提案",
1: "认知提升 • 教育培训 • 成功案例 • 业务增长",
2: "1. 项目目标",
3: "本内容制作项目旨在通过持续的线上宣传、实用教育、增长技巧以及 KiliShop 店主的真实成功案例，增加肯尼亚境内活跃 KiliShop 的数量。内容将帮助受众更容易理解 KiliShop 模式，展示其优势，并鼓励符合条件的潜在店主注册和参与。",
4: "2. 拟制作的内容",
5: "制作教育类视频，讲解 KiliShop 是什么、如何运作、注册要求及运营流程。",
6: "制作认知类内容，突出成为 KiliShop 店主所带来的益处和机会。",
7: "制作 KiliShop 店主成功故事和客户评价，呈现真实经历与实用经验。",
8: "制作增长技巧内容，涵盖包裹领取、客户服务、店铺品牌形象、客户沟通以及常见错误。",
9: "制作幕后花絮（BTS）和失误片段内容，使制作过程更具人情味，并提升受众互动。",
10: "内容发布于 TikTok 和 YouTube；在适当情况下，将精选短视频重新用于其他社交媒体平台。",
11: "3. 内容制作流程",
12: "A. 前期制作",
13: "制定月度内容日历和节目创意。",
14: "提前确定并确认 KiliShop 店铺地点和采访对象。",
15: "研究每个主题，并在拍摄前准备采访问题。",
16: "确认同意及授权发布要求。",
17: "准备并测试摄像机、音频、灯光及其他制作设备。",
18: "为每期节目准备简明制作简报，涵盖制作目标、核心信息、采访对象和预期交付内容。",
19: "B. 现场制作",
20: "根据双方约定的制作时间表，按时到达并完成现场布置。",
21: "进行简短排练或预采访，以建立信任并改善回答效果。",
22: "拍摄清晰的采访画面、辅助 B-roll、KiliShop 运营过程以及相关客户互动（在获得同意的情况下）。",
23: "拍摄幕后花絮和失误片段，用于补充社交媒体内容。",
24: "确保在发布前取得已签署的同意书。",
25: "C. 后期制作",
26: "围绕清晰度、节奏和叙事效果剪辑主视频。",
27: "添加适当的字幕、品牌元素和辅助图形。",
28: "制作 YouTube 吸引人的缩略图及符合平台规范的封面图。",
29: "从合适的长视频内容中制作短视频剪辑版本。",
30: "发布前完成内部准确性和质量审核。",
31: "使用经过优化的标题、描述、文案和行动号召进行发布。",
32: "4. 核心视频主题 / 采访问题",
33: "1. 什么是 KiliShop？它如何运作？",
34: "2. Kilimall 为什么创建 KiliShop？",
35: "3. 成为 KiliShop 店主需要满足哪些要求？注册流程是什么？",
36: "4. KiliShop 店主如何吸引更多包裹领取？",
37: "5. 哪些客户服务做法和店铺品牌建设技巧可以改善客户体验？",
38: "6. 开设 KiliShop 需要哪些文件？需要投入多少资金？",
39: "7. KiliShop 店主常见的错误有哪些？如何避免？",
40: "8. 包裹到达后，应如何通知客户？例如通过电话或消息通知？",
41: "9. 成为 KiliShop 店主的主要益处有哪些？",
42: "10. KiliShop 店主通常会面临哪些挑战？如何克服？",
43: "5. 关键绩效指标（KPI）",
44: "绩效将同时根据内容产出和受众/业务结果进行衡量。以下目标为建议的初始基准；首月获得基准表现数据后，可对其进行调整。",
45: "6. 衡量与归因",
46: "使用 TikTok 和 YouTube 的平台分析数据，跟踪播放量、触达人数、观看时长、留存率、互动量和粉丝增长。",
47: "加入清晰的行动号召，引导感兴趣的观众进入 KiliShop 注册或咨询流程。",
48: "使用简单的潜客跟踪机制（例如活动专属关键词、表单字段、链接或咨询标签），区分由内容带来的咨询和注册。",
49: "维护月度数据看板，对比计划制作内容、已发布内容、受众表现以及可归因的 KiliShop 潜客/注册量。",
50: "复盘前 30–90 天的数据，并根据表现优化主题、内容形式、开场吸引点、发布频率和行动号召。",
51: "7. 质量控制标准",
52: "所有已发布信息必须 100% 准确，并在发布前完成核实。",
53: "教程必须简单、实用且易于理解。",
54: "应尽可能使用真实 KiliShop 案例，以提升可信度。",
55: "所有视频必须通过内部审核后方可发布。",
56: "音频必须清晰，画面应具备充足光线、合理构图并保持稳定。",
57: "信息必须以清晰、吸引人且符合受众理解习惯的方式呈现。",
58: "发布前必须获得采访对象及所有可识别参与者的同意。",
59: "8. 预期成果",
60: "提升受众对 KiliShop 模式的认知和理解。",
61: "建立持续、系统的 KiliShop 教育及推广内容库。",
62: "获得更多来自潜在 KiliShop 店主的有效咨询。",
63: "随着时间推移，增加可归因于内容的 KiliShop 注册量。",
64: "通过成功故事和客户评价，提高现有 KiliShop 店主的曝光度。",
65: "形成可重复使用的视频资产，支持未来的招募、培训和营销活动。",
66: "9. 支持 / 项目团队",
67: "团队：Angela Ajuang，内容制作人；Simon Chege，视频剪辑师",
68: "项目支持：Maryvn Wabwile、James Mburu Wainaina 和 Terry",
}

table_cells = [
    ["KPI", "目标", "衡量方式", "报告频率"],
    ["制作的主视频", "每周 1 条", "已发布且通过审核的视频", "每月"],
    ["短视频/剪辑版视频", "每月 4 条", "已发布的 TikTok/短视频资产", "每月"],
    ["发布一致性", "计划内容的 95%", "已发布内容与计划内容的对比", "每月"],
    ["TikTok 平均播放量", "第 3 个月前达到每条 5,000 次", "TikTok 原生分析数据", "每月"],
    ["YouTube 平均播放量", "第 3 个月前达到每条 1,000 次", "YouTube Analytics", "每月"],
    ["平均互动率", "第 3 个月前达到 5%", "互动量 ÷ 触达人数/播放量 × 100", "每月"],
    ["视频完播率", "第 3 个月前达到 25%", "平台留存率/完播率分析数据", "每月"],
    ["有效 KiliShop 咨询", "第 3 个月前每月 50 条", "可归因于内容的已跟踪咨询", "每月"],
    ["KiliShop 注册/潜客", "第 3 个月前每月 20 个可归因注册", "注册/潜客跟踪", "每月"],
    ["成功故事", "每月 2 个", "已发布的店主评价/案例研究", "每月"],
    ["内容准确性", "100%", "内部事实核查/审批清单", "每条视频"],
    ["内容审批周期", "3 个工作日", "从简报提交到审批完成的跟踪", "每条视频"],
]

doc = Document(src)
for i, text in paragraphs.items():
    p = doc.paragraphs[i]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = text

table = doc.tables[0]
for row, values in zip(table.rows, table_cells):
    for cell, value in zip(row.cells, values):
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].text = value
            for r in cell.paragraphs[0].runs[1:]:
                r.text = ""
        else:
            cell.text = value

# Ensure Chinese glyphs render reliably while preserving the source layout.
for p in doc.paragraphs:
    for run in p.runs:
        run.font.name = "Hiragino Sans GB"
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            from docx.oxml import OxmlElement
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), "Hiragino Sans GB")
        rfonts.set(qn("w:hAnsi"), "Hiragino Sans GB")
        rfonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Hiragino Sans GB"
                    rpr = run._element.get_or_add_rPr()
                    rfonts = rpr.rFonts
                    if rfonts is None:
                        from docx.oxml import OxmlElement
                        rfonts = OxmlElement("w:rFonts")
                        rpr.insert(0, rfonts)
                    rfonts.set(qn("w:ascii"), "Hiragino Sans GB")
                    rfonts.set(qn("w:hAnsi"), "Hiragino Sans GB")
                    rfonts.set(qn("w:eastAsia"), "Hiragino Sans GB")

doc.save(out)
print(out)
